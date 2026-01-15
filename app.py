from fastapi import FastAPI, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from transformers import pipeline
from docx import Document
import tempfile
import mammoth
import PyPDF2
import io

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

summarizer = pipeline(
    "summarization",
    model="facebook/bart-large-cnn"
)

def extract_text(file: UploadFile):
    """Extract text from different file types"""
    content = file.file.read()

    if file.filename.endswith('.txt'):
        return content.decode('utf-8')

    elif file.filename.endswith('.docx'):
        result = mammoth.extract_raw_text(io.BytesIO(content))
        return result.value

    elif file.filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

    else:
        raise ValueError("Unsupported file type")

@app.post("/summarize")
async def summarize(file: UploadFile, length: str = Form("medium")):
    try:
        # Extract text from uploaded file
        text = extract_text(file)
        original_word_count = len(text.split())

        # Limit text length for processing
        text = " ".join(text.split()[:1000])

        # Adjust summary length based on parameter
        length_settings = {
            "short": {"max_length": 80, "min_length": 30},
            "medium": {"max_length": 150, "min_length": 60},
            "detailed": {"max_length": 250, "min_length": 100}
        }

        settings = length_settings.get(length, length_settings["medium"])

        # Generate summary
        summary_result = summarizer(
            text,
            max_length=settings["max_length"],
            min_length=settings["min_length"],
            do_sample=False
        )

        summary_text = summary_result[0]['summary_text']
        summary_word_count = len(summary_text.split())

        # Calculate reduction percentage
        reduction = round((1 - summary_word_count / original_word_count) * 100, 1)

        return {
            "summary": summary_text,
            "stats": {
                "original_words": original_word_count,
                "summary_words": summary_word_count,
                "reduction_percentage": reduction
            }
        }
    except Exception as e:
        return {"error": str(e)}

class DownloadRequest(BaseModel):
    summary: str

@app.post("/download")
async def download(request: DownloadRequest):
    try:
        # Create Word document
        doc = Document()
        doc.add_heading("Document Summary", level=1)
        doc.add_paragraph(request.summary)

        # Save to temporary file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)

        return FileResponse(
            tmp.name,
            filename="summary.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return {"error": str(e)}

@app.get("/")
async def root():
    return {"message": "AI Document Summarizer API is running"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
